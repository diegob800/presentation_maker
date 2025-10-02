import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
import os

url_wiki_r = input("Вставьте свою ссылку: ")  #https://ru.wikipedia.org/wiki/Rust_(игра)

headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
}

document = Document()

style = document.styles['Normal']

def txt_style(style):
    choosen_style=int(input('выбирите номер стиля_ '))###########
    match choosen_style:
        case 1:
            return 'Times New Roman'
        case 2:
            return 'Calibri'
        case 3:
            return 'Arial'
        case _:
            return 'error: недоступный шрифт'
style.font.name = txt_style(style)
def txt_size(style):
    pixel_count=int(input('выбирите номер размера пикселей_ '))
    match pixel_count:
        case 1:
            return '12'
        case 2:
            return '14'
        case 3: 
            return '18'
        case 4:
            return '22'
        case _:
            return '36'
style.font.size = Pt(int(txt_size(style)))

response = requests.get(url_wiki_r, headers=headers)
response.raise_for_status()
soup = BeautifulSoup(response.text, features="html.parser")

headline = soup.find('h1', id="firstHeading").text
print(headline)


document.add_heading(headline, level=1)

content = soup.find('div', {'id': 'mw-content-text'})

if content:
    elements = content.find_all(['h2', 'p'])
    
    for element in elements:
        if element.name == 'h2':
            if element.get_text().strip() in ['Примечания', 'Ссылки', 'Литература']:
                break
            print(f"\n{element.get_text().strip()}")
            print("-" * len(element.get_text().strip()))
            document.add_heading(element.get_text().strip(), level=2)
            
        elif element.name == 'p':
            text = element.get_text().strip()
            if text:
                print(text)
        
                document.add_paragraph(text)
else:
    print("Не удалось найти контент статьи")

save_path = 'C:/Users/Diego/Desktop/NPOEKTbI/presentation_maker/wiki_article.docx'
document.save(save_path)
print(f"\nДокумент сохранен как: {save_path}")
